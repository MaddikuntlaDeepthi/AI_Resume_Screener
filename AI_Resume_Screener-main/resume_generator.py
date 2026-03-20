from docx import Document
import time
import os


def generate_resume(optimized_resume_text):

    folder = "generated_resumes"
    os.makedirs(folder, exist_ok=True)

    doc = Document()

    sections = {
        "summary": [],
        "education": [],
        "projects": [],
        "skills": [],
        "experience": [],
        "certifications": []
    }

    header = []   # ← store name + contact
    current_section = None
    found_first_section = False

    lines = optimized_resume_text.split("\n")

    for line in lines:

        text = line.strip().replace("**", "")

        if text == "":
            continue

        lower = text.lower()

        # Detect sections
        if "summary" in lower:
            current_section = "summary"
            found_first_section = True
            continue

        elif "education" in lower:
            current_section = "education"
            found_first_section = True
            continue

        elif "project" in lower:
            current_section = "projects"
            found_first_section = True
            continue

        elif "skill" in lower:
            current_section = "skills"
            found_first_section = True
            continue

        elif "experience" in lower:
            current_section = "experience"
            found_first_section = True
            continue

        elif "certification" in lower:
            current_section = "certifications"
            found_first_section = True
            continue

        # Store header before first section
        if not found_first_section:
            header.append(text)
        elif current_section:
            sections[current_section].append(text)

    # ----------- WRITE DOCUMENT ------------

    # NAME (first line of header)
    if header:
        doc.add_heading(header[0], level=0)

        for line in header[1:]:
            doc.add_paragraph(line)

    # SUMMARY
    if sections["summary"]:
        doc.add_heading("Professional Summary", level=1)
        for line in sections["summary"]:
            doc.add_paragraph(line)

    # EDUCATION
    if sections["education"]:
        doc.add_heading("Education", level=1)
        for line in sections["education"]:
            doc.add_paragraph(line)

    # PROJECTS
    if sections["projects"]:
        doc.add_heading("Projects", level=1)
        for line in sections["projects"]:
            if line.startswith(("•", "-", "*")):
                doc.add_paragraph(line[1:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(line)

    # SKILLS
    if sections["skills"]:
        doc.add_heading("Skills", level=1)
        for line in sections["skills"]:
            doc.add_paragraph(line.replace("•", "").strip(), style="List Bullet")

    # EXPERIENCE
    if sections["experience"]:
        doc.add_heading("Experience", level=1)
        for line in sections["experience"]:
            doc.add_paragraph(line)

    # CERTIFICATIONS
    if sections["certifications"]:
        doc.add_heading("Certifications", level=1)
        for line in sections["certifications"]:
            doc.add_paragraph(line)

    file_name = f"optimized_resume_{int(time.time())}.docx"
    file_path = os.path.join(folder, file_name)

    doc.save(file_path)

    return file_path